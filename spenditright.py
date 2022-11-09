import sys
import sqlite3
import datetime

from docx import Document
from docx.shared import Inches
from PyQt5 import uic, QtWidgets, QtCore, QtGui
from PyQt5.Qt import *
from PyQt5.QtWidgets import *
from PyQt5.QtChart import QChart, QChartView, QPieSeries, QPieSlice, QBarSeries, QBarSet, QBarCategoryAxis
from PyQt5.QtGui import QColor, QIcon

month_names = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май',
               'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь',
               'Ноябрь', 'Декабрь']

month = month_names[datetime.datetime.now().month - 1]

DATABASE = sqlite3.connect('database.sqlite')
cur = DATABASE.cursor()

# если наступил новый месяц, добавляем его в months.txt
with open('months.txt', encoding='utf-8', mode='r') as file:
    MONTHS = list(map(lambda x: x.replace('\n', ''), file.readlines()))

if month_names[datetime.datetime.now().month - 1] not in MONTHS:
    MONTHS.append(month)
    with open('months.txt', encoding='utf-8', mode='w') as file:
        file.write(month)
    # также добавляем поле с месяцем в salary.sqlite
    with sqlite3.connect('salary.sqlite') as salary_db:
        salary_cur = salary_db.cursor()
        salary_cur.execute("""INSERT INTO months(month, salary) VALUES(?, 0)""", (month, ))
        salary_db.commit()

with sqlite3.connect('salary.sqlite') as salary_db:
    salary_cur = salary_db.cursor()
    SALARY = salary_cur.execute("""SELECT salary FROM months WHERE month = ?""", (month, )).fetchall()[0][0]


class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui/main_window.ui', self)

        self.sum_ = 0
        self.items = []
        self.slices = []

        self.series = QPieSeries(self)
        self.series.setHoleSize(0.50)

        self.initUI()

    def initUI(self):
        self.setStyleSheet("background-color: white;")

        self.month = QLabel(MONTHS[-1], self)
        self.month.resize(200, 50)
        self.month.setStyleSheet("font-size: 26px;")
        self.month.move(950, 20)

        self.slices = []

        # получение категорий за действующий месяц из базы данных
        categories = {i[0]: 0 for i in cur.execute("""SELECT Categories.name FROM Categories
                                                      INNER JOIN Months ON Months.category = Categories.id
                                                      WHERE Months.name = ?""",
                                                   (self.month.text(),)).fetchall()}
        # создание слайсов для диаграммы
        for key, val in categories.items():
            for i in cur.execute("""SELECT Purchases.price FROM Purchases
                                    INNER JOIN Months ON Months.category = Categories.id
                                    INNER JOIN Categories ON Categories.purchase = Purchases.id
                                    WHERE Categories.name = ? AND Months.name = ?""",
                                 (key, self.month.text(),)).fetchall():
                categories[key] += i[0]
            self.slices.append(QPieSlice(key, categories[key]))

        for i in self.slices:
            self.series.append(i)

        # создание графика
        self.chart = QChart()
        self.chart.legend().hide()
        self.chart.addSeries(self.series)
        self.chart.setAnimationOptions(QChart.SeriesAnimations)
        self.chart.setBackgroundVisible(False)

        # отображение графика
        self.chart_display = QChartView(self.chart, self)
        self.chart_display.setRenderHint(QPainter.Antialiasing)
        self.chart_display.resize(600, 600)
        self.chart_display.move(700, 50)
        self.series.hovered.connect(self.slice_hovered)
        self.chart_display.setStyleSheet('background: transparent;')

        self.saved_lbl = QLabel(self)
        self.saved_lbl.setStyleSheet('font-size: 24px; background: transparent;')
        self.saved_lbl.resize(300, 100)
        self.saved_lbl.move(890, 550)

        self.expenses = QLabel("Расходы:", self)
        self.expenses.resize(100, 20)
        self.expenses.setStyleSheet("font-size: 20px; background: transparent;")
        self.expenses.move(957, 90)

        self.label_in_chart = QLabel(self)
        self.label_in_chart.resize(150, 50)
        self.label_in_chart.setStyleSheet('font-size: 20px;')
        self.label_in_chart.move(940, 320)
        self.view_items_in_list_widget()

        # кастомизация скроллбара
        scroll_bar = QScrollBar(self)
        scroll_bar.setStyleSheet("""QScrollBar:vertical {              
                                    border: none;
                                    background: white;
                                    width: 5px;               
                                    margin: 0px 0px 0px 0px;
                                    }
                                    QScrollBar::handle:vertical {
                                    background: black;
                                    min-height: 0px;
                                    }""")
        self.listWidget.setVerticalScrollBar(scroll_bar)

        self.main_window_button.clicked.connect(self.change_tab)
        self.comparison_window_button.clicked.connect(self.change_tab)
        self.settings_button.clicked.connect(self.change_tab)

        self.hovered_slice_category = QLabel(self)
        self.hovered_slice_category.setText("Выбранная категория")
        self.hovered_slice_category.resize(170, 20)
        self.hovered_slice_category.move(950, 380)

        self.addCategoryButton.clicked.connect(self.add_category)
        self.deleteCategoryButton.clicked.connect(self.delete_category)

        self.previos_month_btn = QPushButton(self)
        self.previos_month_btn.resize(40, 40)
        self.previos_month_btn.move(880, 29)
        self.set_buttons_icon(self.previos_month_btn, 'images/left.png')
        self.previos_month_btn.setStyleSheet('border: 1px solid black; background-color: white; border-radius: 5px;')
        self.previos_month_btn.clicked.connect(self.set_month)

        self.next_month_btn = QPushButton(self)
        self.next_month_btn.resize(40, 40)
        self.next_month_btn.move(1080, 29)
        self.set_buttons_icon(self.next_month_btn, 'images/right.png')
        self.next_month_btn.setStyleSheet('border: 1px solid black; background-color: white; border-radius: 5px;')
        self.next_month_btn.clicked.connect(self.set_month)

        self.monthly_report.clicked.connect(self.save_report)

    def slice_hovered(self, slice: QPieSlice, state) -> None:

        """Анимация при наведении на слайс"""

        if state:
            slice.setExploded(True)
            self.hovered_slice_category.setText(slice.label() + ' ' + str(round(slice.value())) + 'руб.')
        else:
            slice.setExploded(False)
            slice.setLabelVisible(False)
            self.hovered_slice_category.setText("Выбранная категория")

    def set_button_labels_visible(self) -> None:

        """Добавление списка с покупками под активную кнопку"""

        index = 0
        for i in range(len(self.slices)):
            if self.slices[i].label() == self.sender().text():
                index = i
                break
        if index < len(self.slices):
            self.refresh_items_in_list_widget()

    @staticmethod
    def set_buttons_icon(object: QPushButton, file_name: str) -> None:

        """Изменение иконки у нажатой кнопки"""

        object.setIcon(QIcon(file_name))

    def view_items_in_list_widget(self) -> None:

        """Добавлие категорий в список"""

        self.sum_ = 0
        self.items = []

        for i in self.slices:
            self.sum_ += round(i.value())

            # создание элемента ListWidget для дальнейшего преобразования его в кнопку
            element = QtWidgets.QListWidgetItem()
            element.setSizeHint(QtCore.QSize(100, 50))
            self.listWidget.addItem(element)

            # создаем кнопку
            btn = QPushButton(self)
            btn.setText(str(i.label()))
            btn.setStyleSheet(f"border-radius: 0;"
                              f"border-bottom: 1px solid #c4c4c4;"
                              f"border-top: none;"
                              f"border-right: none;"
                              f"border-left: none;")
            btn.setIcon(QIcon("images/right.png"))
            btn.clicked.connect(self.set_button_labels_visible)

            self.items.append(btn)
            self.listWidget.setItemWidget(element, btn)
            self.listWidget.scrollToItem(element)
        self.label_in_chart.setText("    Итого:\n" + str(self.sum_) + " руб.")
        self.set_saved_value()

    def refresh_items_in_list_widget(self) -> None:

        """Функция добавляет под нажатую кнопку список с покупками и две функциональные кнопки,
         а под неактивной кнопкой убирает их"""

        # оставляем в списке категорий только кнопки с названиями категорий
        self.items = list(filter(lambda x: isinstance(x, QPushButton) and x.text() != '+' and x.text() != '-',
                                 self.items))

        index = [i for i in range(len(self.items)) if self.items[i].text() == self.sender().text()][0]

        temp_list = self.items
        self.items = []

        self.listWidget.clear()

        for i in range(len(temp_list)):
            element = QtWidgets.QListWidgetItem()
            element.setSizeHint(QtCore.QSize(100, 50))
            self.listWidget.addItem(element)

            btn = QPushButton(self)
            btn.setText(str(temp_list[i].text()))
            btn.setStyleSheet(f"border-radius: 0;"
                              f"border-bottom: 1px solid #c4c4c4;"
                              f"border-top: none;"
                              f"border-right: none;"
                              f"border-left: none;")
            btn.setIcon(QIcon("images/right.png"))
            if i != index:
                btn.clicked.connect(self.set_button_labels_visible)
                self.items.append(btn)
                self.listWidget.setItemWidget(element, btn)
            else:
                self.set_buttons_icon(btn, "images/down.png")
                btn.clicked.connect(self.set_button_labels_visible)
                self.items.append(btn)
                self.listWidget.setItemWidget(element, btn)

                element = QtWidgets.QListWidgetItem()
                element.setSizeHint(QtCore.QSize(100, 100))
                self.listWidget.addItem(element)

                lbl = QTextEdit(self)
                lbl.setDisabled(True)
                self.listWidget.setItemWidget(element, lbl)

                # достаем покупки и их цены из базы данных и заносим в список
                for n, elem in enumerate(cur.execute("""SELECT Purchases.name, Purchases.price FROM Purchases
                                                        JOIN Categories ON Categories.purchase = Purchases.id
                                                        INNER JOIN Months ON Months.category = Categories.id
                                                        WHERE Categories.name = ? AND Months.name = ?""",
                                                     (self.items[-1].text(), self.month.text(),)).fetchall()):
                    lbl.append(f'{n + 1}. {elem[0]} - {elem[1]} руб.')
                self.items.append(lbl)

                # добавляем кнопки для добавления/удаления покупок
                for j in range(2):
                    element = QtWidgets.QListWidgetItem()
                    element.setSizeHint(QtCore.QSize(100, 30))
                    self.listWidget.addItem(element)

                    btn = QPushButton(self)
                    btn.setText(['+', '-'][j])
                    if btn.text() == '+':
                        btn.clicked.connect(self.add_purchase)
                    else:
                        btn.clicked.connect(self.delete_purchase)
                    btn.setStyleSheet(f"border-radius: 5;"
                                      f"border: 1px solid #c4c4c4;"
                                      )
                    self.listWidget.setItemWidget(element, btn)
                    self.items.append(btn)
        self.set_saved_value()

    def set_month(self) -> None:

        """Функция меняет месяц и обновляет информацию по действующему месяцу"""

        if self.sender().pos().x() == 880:
            self.month.setText(MONTHS[(MONTHS.index(self.month.text()) - 1) % len(MONTHS)])
        else:
            self.month.setText(MONTHS[(MONTHS.index(self.month.text()) + 1) % len(MONTHS)])
        self.listWidget.clear()
        self.slices.clear()
        self.series.clear()
        categories = {i[0]: 0 for i in cur.execute("""SELECT Categories.name FROM Categories
                                                      INNER JOIN Months ON Months.category = Categories.id
                                                      WHERE Months.name = ?""",
                                                   (self.month.text(),)).fetchall()}
        for key, val in categories.items():
            for i in cur.execute("""SELECT Purchases.price FROM Purchases
                                    INNER JOIN Months ON Months.category = Categories.id
                                    INNER JOIN Categories ON Categories.purchase = Purchases.id
                                    WHERE Categories.name = ? AND Months.name = ?""",
                                 (key, self.month.text(),)).fetchall():
                categories[key] += i[0]
            self.slices.append(QPieSlice(key, categories[key]))
        for i in self.slices:
            self.series.append(i)
        self.view_items_in_list_widget()

    def add_category(self) -> None:

        """Функция запрашивает название категории у пользователя и добавляет ее в список"""

        text, ok = QInputDialog.getText(self, 'Добавить категорию', 'Введите название категории:')
        if ok and text != '+' and text != '-':
            self.items = list(filter(lambda x: isinstance(x, QPushButton) and x.text() != '+' and x.text() != '-',
                                     self.items))  # оставляем в items только кнопки с названиями категорий
            if text.lower() not in [i.text().lower() for i in self.items]:
                self.slices.append(QPieSlice(text, 0.1))
                for i in self.slices:
                    self.series.append(i)
                self.listWidget.clear()
                self.view_items_in_list_widget()
                self.set_saved_value()

    def delete_category(self) -> None:

        """Запрашивает название категории у пользователя и удаляет ее из списка"""

        text, ok = QInputDialog.getText(self, 'Удалить категорию',
                                        'Введите название категории, которую хотите удалить:')
        if ok:
            pie_slice = None
            categories_name = None
            for i in range(len(self.slices)):
                if self.slices[i].label() == text:
                    pie_slice = self.slices[i]
                    categories_name = text
                    del self.slices[i]
                    break

            cur.execute("""DELETE from Months
                            WHERE Months.name = ?
                            AND
                            Months.category IN (SELECT id FROM Categories
                                                WHERE name = ?)""",
                        (self.month.text(), categories_name,)).fetchall()
            DATABASE.commit()

            cur.execute("""DELETE from Categories
                           WHERE Categories.id NOT IN (SELECT category FROM Months)""").fetchall()
            DATABASE.commit()

            cur.execute("""DELETE from Purchases
                           WHERE Purchases.id NOT IN (SELECT purchase FROM Categories)""").fetchall()
            DATABASE.commit()

            self.series.remove(pie_slice)
            self.listWidget.clear()
            self.view_items_in_list_widget()
            self.set_saved_value()
            if isinstance(window.tabwidget.widget(1), Analys):
                window.tabwidget.removeTab(1)
                window.tabwidget.insertTab(1, Analys(), "Tab 2")
            else:
                window.tabwidget.removeTab(1)
                window.tabwidget.insertTab(1, DetailedAnalys(), "Tab 2")

    def add_purchase(self) -> None:

        """Добавляет покупку под выбранную категорию"""

        purchase_name, ok1 = QInputDialog.getText(self, 'Добавить покупку',
                                                  'Введите название покупки:')
        if ok1:
            purchase_price, ok2 = QInputDialog.getText(self, 'Добавить покупку',
                                                       'Введите цену:')
            if ok2 and purchase_price.isdigit():
                index = 0
                for i in range(len(self.items)):
                    if isinstance(self.items[i], QTextEdit):
                        index = i
                        break
                text = list(filter(lambda x: x != '', self.items[index].toPlainText().split("\n")))  # достаем из PlainText все покупки и сразу убираем символ переноса строки
                self.items[index].append(f'{len(text) + 1}. {purchase_name} - {purchase_price} руб.')
                for i in range(len(self.series.slices())):
                    if self.series.slices()[i].label() == self.items[index - 1].text():
                        self.series.slices()[i].setValue(self.series.slices()[i].value() + float(purchase_price))
                        self.sum_ += float(purchase_price)
                self.label_in_chart.setText("    Итого:\n" + str(self.sum_) + " руб.")
                self.set_saved_value()

                cur.execute("""INSERT INTO purchases(name, price) VALUES(?, ?)""",
                            (purchase_name, float(purchase_price),))
                DATABASE.commit()
                pur_index = max(cur.execute("""SELECT id FROM purchases""").fetchall())
                cur.execute("""INSERT INTO categories(name, purchase) VALUES (?, ?)""",
                            (self.items[index - 1].text(), pur_index[0],))
                DATABASE.commit()
                cat_index = max(cur.execute("""SELECT id FROM categories""").fetchall())
                cur.execute("""INSERT INTO months(name, category) VALUES (?, ?)""",
                            (self.month.text(), cat_index[0],))
                DATABASE.commit()
                if isinstance(window.tabwidget.widget(1), Analys):
                    window.tabwidget.removeTab(1)
                    window.tabwidget.insertTab(1, Analys(), "Tab 2")
                else:
                    window.tabwidget.removeTab(1)
                    window.tabwidget.insertTab(1, DetailedAnalys(), "Tab 2")

    def delete_purchase(self):

        """Удаляет покупку из выбранной категории"""

        purchase_index, ok = QInputDialog.getText(self, 'Удалить покупку',
                                                  'Введите номер покупки:')
        if ok and purchase_index.isdigit():
            index = 0
            for i in range(len(self.items)):
                if isinstance(self.items[i], QTextEdit):
                    index = i
                    break
            category = self.items[index - 1].text()
            text = list(filter(lambda x: x != '', self.items[index].toPlainText().split("\n")))[int(purchase_index) - 1]
            purchase_name, purchase_price = text.split()[1], int(text.split()[3])

            cur.execute("""DELETE from Months
                           WHERE Months.name = ? 
                           AND Months.category = (SELECT id FROM Categories
                                                    WHERE Categories.name = ? 
                                                    AND
                                                    Categories.purchase = (SELECT id FROM Purchases
                                                                           WHERE Purchases.name = ?
                                                                           AND 
                                                                           Purchases.price = ?))""",
                        (self.month.text(), category, purchase_name, purchase_price)).fetchall()
            DATABASE.commit()
            cur.execute("""DELETE from Categories
                            WHERE Categories.name = ? 
                            AND Categories.purchase = (SELECT id FROM Purchases
                                                       WHERE Purchases.name = ? 
                                                       AND 
                                                       Purchases.price = ?)
                                                       AND 
                                                       Categories.id NOT IN (SELECT category FROM Months)""",
                        (category, purchase_name, purchase_price)).fetchall()
            DATABASE.commit()
            cur.execute("""DELETE from Purchases
                           WHERE Purchases.name = ? 
                           AND 
                           Purchases.price = ? 
                           AND 
                           Purchases.id NOT IN (SELECT purchase FROM Categories)""",
                        (purchase_name, purchase_price)).fetchall()
            DATABASE.commit()

            self.items[index].clear()
            for n, elem in enumerate(cur.execute("""SELECT Purchases.name, Purchases.price FROM Purchases
                                                    JOIN Categories ON Categories.purchase = Purchases.id
                                                    INNER JOIN Months ON Months.category = Categories.id
                                                    WHERE Categories.name = ? AND Months.name = ?""",
                                                 (category, self.month.text(),)).fetchall()):
                self.items[index].append(f'{n + 1}. {elem[0]} - {elem[1]} руб.')
            self.sum_ -= purchase_price
            self.label_in_chart.setText("    Итого:\n" + str(self.sum_) + " руб.")
            self.set_saved_value()
            if isinstance(window.tabwidget.widget(1), Analys):
                window.tabwidget.removeTab(1)
                window.tabwidget.insertTab(1, Analys(), "Tab 2")
            else:
                window.tabwidget.removeTab(1)
                window.tabwidget.insertTab(1, DetailedAnalys(), "Tab 2")

    def set_saved_value(self) -> None:

        """Расчет сэкономленных денег в процентах"""

        global SALARY

        saved = (SALARY - self.sum_) / SALARY * 100 if SALARY != 0 else 0
        if saved < 0:
            saved = 0
        self.saved_lbl.setText(f'Сэкономлено: {round(float(saved), 2)}%')

    def save_report(self) -> None:

        """Составление отчета за месяц в формате docx"""

        path = QFileDialog.getExistingDirectory(self, "Выбрать папку", ".")
        document = Document()
        document.add_heading(f'{self.month.text()} {datetime.datetime.now().year}', 0)
        for i in self.slices:
            document.add_heading(i.label(), level=1)
            for j, elem in enumerate(cur.execute("""SELECT Purchases.name, Purchases.price FROM Purchases
                                                    JOIN Categories ON Categories.purchase = Purchases.id
                                                    INNER JOIN Months ON Months.category = Categories.id
                                                    WHERE Categories.name = ? AND Months.name = ?""",
                                                 (i.label(), self.month.text(),)).fetchall()):
                document.add_paragraph(f'{j + 1}. {elem[0]} - {elem[1]} руб.')
        document.add_paragraph()
        document.add_paragraph(f'Доход: {SALARY} руб.')
        document.add_paragraph(f'Итого: {self.sum_} руб.')
        document.add_paragraph(self.saved_lbl.text())
        document.save(path + '/' + f'Отчет {self.month.text()}.docx')
        print('Отчет успешно сохранен!')

    def change_tab(self) -> None:

        """Смена вкладки"""

        window.tabwidget.setCurrentIndex(["Главное окно", "Сравнение", "Настройки"].index(self.sender().text()))
        window.setWindowTitle(f'Spend it right | {self.sender().text()}')


class Analys(QWidget):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui/default_analys.ui', self)
        self.series1 = QPieSeries(self)
        self.series1.setHoleSize(0.50)
        self.series2 = QPieSeries(self)
        self.series2.setHoleSize(0.50)
        self.slices1 = []
        self.slices2 = []
        self.initUI()

    def initUI(self):
        self.month1.setText(MONTHS[-1])
        self.month2.setText(MONTHS[-1])

        self.main_window_button.clicked.connect(self.change_tab)
        self.comparison_window_button.clicked.connect(self.change_tab)
        self.settings_button.clicked.connect(self.change_tab)

        categories = {i[0]: 0 for i in cur.execute("""SELECT Categories.name FROM Categories
                                                      INNER JOIN Months ON Months.category = Categories.id
                                                      WHERE Months.name = ?""",
                                                   (self.month1.text(),)).fetchall()}
        for key, val in categories.items():
            for i in cur.execute("""SELECT Purchases.price FROM Purchases
                                    INNER JOIN Months ON Months.category = Categories.id
                                    INNER JOIN Categories ON Categories.purchase = Purchases.id
                                    WHERE Categories.name = ? AND Months.name = ?""",
                                 (key, self.month1.text(),)).fetchall():
                categories[key] += i[0]
            self.slices1.append(QPieSlice(key, categories[key]))

        for i in self.slices1:
            self.series1.append(i)

        categories = {i[0]: 0 for i in cur.execute("""SELECT Categories.name FROM Categories
                                                      INNER JOIN Months ON Months.category = Categories.id
                                                      WHERE Months.name = ?""",
                                                   (self.month2.text(),)).fetchall()}
        for key, val in categories.items():
            for i in cur.execute("""SELECT Purchases.price FROM Purchases
                                    INNER JOIN Months ON Months.category = Categories.id
                                    INNER JOIN Categories ON Categories.purchase = Purchases.id
                                    WHERE Categories.name = ? AND Months.name = ?""",
                                 (key, self.month2.text(),)).fetchall():
                categories[key] += i[0]
            self.slices2.append(QPieSlice(key, categories[key]))

        for i in self.slices2:
            self.series2.append(i)

        self.set_saved_value(0)
        self.set_saved_value(1)

        # создание первого графика
        self.chart1 = QChart()
        self.chart1.legend().hide()
        self.chart1.addSeries(self.series1)
        self.chart1.setAnimationOptions(QChart.SeriesAnimations)
        self.chart1.setBackgroundVisible(False)

        # визуалицация первого графика
        self.chart_display1 = QChartView(self.chart1, self)
        self.chart_display1.setRenderHint(QPainter.Antialiasing)
        self.chart_display1.resize(400, 400)
        self.chart_display1.move(300, 100)
        self.series1.hovered.connect(self.slice_hovered)
        self.chart_display1.setStyleSheet('background: transparent;')

        self.month1_left_btn.setIcon(QIcon('images/left.png'))
        self.month1_left_btn.clicked.connect(self.set_month)
        self.month1_right_btn.setIcon(QIcon('images/right.png'))
        self.month1_right_btn.clicked.connect(self.set_month)

        # создание второго графика
        self.chart2 = QChart()
        self.chart2.legend().hide()
        self.chart2.addSeries(self.series2)
        self.chart2.setAnimationOptions(QChart.SeriesAnimations)
        self.chart2.setBackgroundVisible(False)

        # визуалицация первого графика
        self.chart_display2 = QChartView(self.chart2, self)
        self.chart_display2.setRenderHint(QPainter.Antialiasing)
        self.chart_display2.resize(400, 400)
        self.chart_display2.move(800, 100)
        self.series2.hovered.connect(self.slice_hovered)
        self.chart_display2.setStyleSheet('background: transparent;')

        self.month2_left_btn.setIcon(QIcon('images/left.png'))
        self.month2_left_btn.clicked.connect(self.set_month)
        self.month2_right_btn.setIcon(QIcon('images/right.png'))
        self.month2_right_btn.clicked.connect(self.set_month)

    @staticmethod
    def slice_hovered(slice: QPieSlice, state) -> None:

        """Анимация при наведении на слайс"""

        if state:
            slice.setExploded(True)
            slice.setLabelVisible(True)
        else:
            slice.setExploded(False)
            slice.setLabelVisible(False)

    def set_month(self) -> None:

        """Смена месяца и обновление графика"""

        if self.sender().pos().x() == 400:
            self.month1.setText(MONTHS[(MONTHS.index(self.month1.text()) - 1) % len(MONTHS)])
        elif self.sender().pos().x() == 570:
            self.month1.setText(MONTHS[(MONTHS.index(self.month1.text()) + 1) % len(MONTHS)])
        elif self.sender().pos().x() == 900:
            self.month1.setText(MONTHS[(MONTHS.index(self.month1.text()) - 1) % len(MONTHS)])
        else:
            self.month1.setText(MONTHS[(MONTHS.index(self.month1.text()) + 1) % len(MONTHS)])
        if self.sender().pos().x() in [400, 570]:
            self.slices1.clear()
            self.slices1.clear()
            categories = {i[0]: 0 for i in cur.execute("""SELECT Categories.name FROM Categories
                                                          INNER JOIN Months ON Months.category = Categories.id
                                                          WHERE Months.name = ?""",
                                                       (self.month1.text(),)).fetchall()}
            for key, val in categories.items():
                for i in cur.execute("""SELECT Purchases.price FROM Purchases
                                        INNER JOIN Months ON Months.category = Categories.id
                                        INNER JOIN Categories ON Categories.purchase = Purchases.id
                                        WHERE Categories.name = ? AND Months.name = ?""",
                                     (key, self.month1.text(),)).fetchall():
                    categories[key] += i[0]
                self.slices1.append(QPieSlice(key, categories[key]))
            self.series1.clear()
            for i in self.slices1:
                self.series1.append(i)
            self.set_saved_value(0)
        else:
            self.slices2.clear()
            self.slices2.clear()
            categories = {i[0]: 0 for i in cur.execute("""SELECT Categories.name FROM Categories
                                                          INNER JOIN Months ON Months.category = Categories.id
                                                          WHERE Months.name = ?""",
                                                       (self.month2.text(),)).fetchall()}
            for key, val in categories.items():
                for i in cur.execute("""SELECT Purchases.price FROM Purchases
                                        INNER JOIN Months ON Months.category = Categories.id
                                        INNER JOIN Categories ON Categories.purchase = Purchases.id
                                        WHERE Categories.name = ? AND Months.name = ?""",
                                     (key, self.month2.text(),)).fetchall():
                    categories[key] += i[0]
                self.slices2.append(QPieSlice(key, categories[key]))
            self.series2.clear()
            for i in self.slices2:
                self.series2.append(i)
            self.set_saved_value(1)

    def set_saved_value(self, chart_index) -> None:

        """Получение суммы сэкономленных средтств в процентах"""

        sum_ = 0
        for i in [self.slices1, self.slices2][chart_index]:
            sum_ += round(i.value())
        saved = (SALARY - sum_) / SALARY * 100 if SALARY != 0 else 0
        if saved < 0:
            saved = 0
        [self.saved1, self.saved2][chart_index].setText(f'Сэкономлено: {round(float(saved), 2)}%')

    def change_tab(self) -> None:

        """Смена вкладки"""

        window.tabwidget.setCurrentIndex(["Главное окно", "Сравнение", "Настройки"].index(self.sender().text()))
        window.setWindowTitle(f'Spend it right | {self.sender().text()}')


class DetailedAnalys(QWidget):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui/detailed_analys.ui', self)
        self.series = QBarSeries()
        self.initUI()

    def initUI(self):
        self.month1.setText(MONTHS[-1])
        self.month2.setText(MONTHS[-1])

        self.main_window_button.clicked.connect(self.change_tab)
        self.comparison_window_button.clicked.connect(self.change_tab)
        self.settings_button.clicked.connect(self.change_tab)

        # создание столбчатой диаграммы
        self.chart = QChart()
        self.chart.legend().hide()
        self.chart.addSeries(self.series)
        self.chart.setAnimationOptions(QChart.SeriesAnimations)
        self.chart.setBackgroundVisible(False)

        self.chart.setTheme(
            window.tabwidget.widget(1).themes[window.tabwidget.widget(1).themes_names.index(
                window.tabwidget.widget(1).set_theme_combo_box.currentText())])

        self.month1_left_btn.setIcon(QIcon('images/left.png'))
        self.month1_left_btn.clicked.connect(self.set_month)
        self.month1_right_btn.setIcon(QIcon('images/right.png'))
        self.month1_right_btn.clicked.connect(self.set_month)

        self.month2_left_btn.setIcon(QIcon('images/left.png'))
        self.month2_left_btn.clicked.connect(self.set_month)
        self.month2_right_btn.setIcon(QIcon('images/right.png'))
        self.month2_right_btn.clicked.connect(self.set_month)

        # визуализация столбчатой диаграммы
        self.chart_view = QChartView(self.chart, self)
        self.chart_view.setRenderHint(QPainter.Antialiasing)
        self.chart_view.setStyleSheet(f'background: transparent;')

        self.chart_view.move(300, 20)
        self.chart_view.resize(880, 500)

        self.load_info()

    def set_month(self) -> None:

        """Смена месяца"""

        self.chart.removeAxis(self.chart.axes()[0])
        if self.sender().pos().x() == 390:
            self.month1.setText(MONTHS[(MONTHS.index(self.month1.text()) - 1) % len(MONTHS)])
        elif self.sender().pos().x() == 560:
            self.month1.setText(MONTHS[(MONTHS.index(self.month1.text()) + 1) % len(MONTHS)])
        elif self.sender().pos().x() == 890:
            self.month1.setText(MONTHS[(MONTHS.index(self.month1.text()) - 1) % len(MONTHS)])
        else:
            self.month1.setText(MONTHS[(MONTHS.index(self.month1.text()) + 1) % len(MONTHS)])
        self.load_info()

    def load_info(self):

        """Обновление графика"""

        self.categories_and_sums_1 = {i[0]: 0 for i in cur.execute("""SELECT Categories.name FROM Categories
                                                                          INNER JOIN Months ON Months.category = Categories.id
                                                                          WHERE Months.name = ?""",
                                                                   (self.month1.text(),)).fetchall()}
        self.categories_and_sums_2 = {i[0]: 0 for i in cur.execute("""SELECT Categories.name FROM Categories
                                                                          INNER JOIN Months ON Months.category = Categories.id
                                                                          WHERE Months.name = ?""",
                                                                   (self.month2.text(),)).fetchall()}
        for i in cur.execute("""SELECT Categories.name, Purchases.price FROM Categories, Purchases
                                        INNER JOIN Months ON Months.category = Categories.id
                                        WHERE Months.name = ? AND Categories.purchase = Purchases.id""",
                             (self.month1.text(),)).fetchall():
            self.categories_and_sums_1[i[0]] += i[1]
        for i in cur.execute("""SELECT Categories.name, Purchases.price FROM Categories, Purchases
                                        INNER JOIN Months ON Months.category = Categories.id
                                        WHERE Months.name = ? AND Categories.purchase = Purchases.id""",
                             (self.month2.text(),)).fetchall():
            self.categories_and_sums_2[i[0]] += i[1]

        self.list_of_sums_1 = []
        self.list_of_sums_2 = []

        for i in sorted(list(set(self.categories_and_sums_1.keys()).union(set(self.categories_and_sums_2.keys())))):
            if i in self.categories_and_sums_1 and i in self.categories_and_sums_2:
                self.list_of_sums_1.append(self.categories_and_sums_1[i])
                self.list_of_sums_2.append(self.categories_and_sums_2[i])
            elif i in self.categories_and_sums_1 and i not in self.categories_and_sums_2:
                self.list_of_sums_1.append(self.categories_and_sums_1[i])
                self.list_of_sums_2.append(0)
            else:
                self.list_of_sums_1.append(0)
                self.list_of_sums_2.append(self.categories_and_sums_2[i])

        self.bars = [QBarSet(''), QBarSet('')]

        self.bars[0].append(self.list_of_sums_1)
        self.bars[1].append(self.list_of_sums_2)

        self.chart.removeSeries(self.series)

        self.series = QBarSeries()
        for i in self.bars:
            self.series.append(i)
        self.chart.addSeries(self.series)

        self.axisX = QBarCategoryAxis()
        for i in sorted(list(set(self.categories_and_sums_1.keys()).union(set(self.categories_and_sums_2.keys())))):
            self.axisX.append(i)

        self.chart.addAxis(self.axisX, Qt.AlignBottom)

    def change_tab(self) -> None:

        """Смена вкладки"""

        window.tabwidget.setCurrentIndex(["Главное окно", "Сравнение", "Настройки"].index(self.sender().text()))
        window.setWindowTitle(f'Spend it right | {self.sender().text()}')


class Settings(QWidget):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui/settings.ui', self)

        self.themes_names = ['Light Theme', 'Toxic Theme', 'Yellow Theme', 'Brown Sand', 'White and black', 'Qt Theme']
        self.themes = [QChart.ChartThemeLight, QChart.ChartThemeDark, QChart.ChartThemeBlueCerulean,
                       QChart.ChartThemeBrownSand, QChart.ChartThemeHighContrast, QChart.ChartThemeQt]

        self.initUI()

    def initUI(self):
        self.main_window_button.clicked.connect(self.change_tab)
        self.comparison_window_button.clicked.connect(self.change_tab)
        self.settings_button.clicked.connect(self.change_tab)

        self.detailed_analys_ind.setStyleSheet('''QCheckBox::indicator {
                                                    width:  33px;
                                                    height: 33px;
                                                }
                                                ''')
        self.detailed_analys_ind.clicked.connect(self.change_analys_mode)

        for i in self.themes_names:
            self.set_theme_combo_box.addItem(i)

        self.set_theme_btn.clicked.connect(self.change_chart_theme)
        self.set_salary_btn.clicked.connect(self.set_salary)
        self.salary_spin_box.setValue(SALARY)

    def change_analys_mode(self) -> None:

        """Смена режима анализа"""

        window.tabwidget.removeTab(1)
        window.tabwidget.insertTab(1, DetailedAnalys() if self.sender().isChecked() else Analys(), "Tab 2")

    def change_chart_theme(self) -> None:

        """Смена темы в диаграммах"""
        
        index = self.themes_names.index(self.set_theme_combo_box.currentText())

        window.tabwidget.widget(0).chart.setTheme(self.themes[index])
        if isinstance(window.tabwidget.widget(1), Analys):
            window.tabwidget.widget(1).chart1.setTheme(self.themes[index])
            window.tabwidget.widget(1).chart2.setTheme(self.themes[index])
        else:
            window.tabwidget.widget(1).chart.setTheme(self.themes[index])

    def set_salary(self) -> None:

        """Смена значения дохода"""

        global SALARY
        with sqlite3.connect('salary.sqlite') as salary_db2:
            salary_cur2 = salary_db2.cursor()
            salary_cur2.execute("""UPDATE months
                                   SET salary = ?
                                   WHERE month = ?""", (self.salary_spin_box.value(), month))
            salary_db2.commit()
            SALARY = salary_cur2.execute("""SELECT salary FROM months WHERE month = ?""", (month, )).fetchall()[0][0]

        window.tabwidget.removeTab(0)
        window.tabwidget.insertTab(0, MainWindow(), 'Tab 1')
        window.tabwidget.removeTab(1)
        window.tabwidget.insertTab(1, Analys(), 'Tab 2')


    def change_tab(self) -> None:

        """Смена вкладки"""

        window.tabwidget.setCurrentIndex(["Главное окно", "Сравнение", "Настройки"].index(self.sender().text()))
        window.setWindowTitle(f'Spend it right | {self.sender().text()}')


class TabWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowIcon(QIcon("images/window_icon.png"))
        self.setFixedSize(1280, 720)
        self.setStyleSheet("background-color: white;")
        self.tabwidget = QTabWidget(self)
        self.tabwidget.resize(1280, 720)
        self.tabwidget.setStyleSheet("background-color: white;")
        self.setWindowTitle("Spend it right | Главное окно")

        self.tabwidget.addTab(MainWindow(), "Tab 1")
        self.tabwidget.addTab(Analys(), "Tab 2")
        self.tabwidget.addTab(Settings(), "Tab 3")


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = TabWidget()
    window.tabwidget.setStyleSheet("QTabBar:tab{background-color: transparent; color: transparent;}")
    window.tabwidget.move(0, -30)
    window.tabwidget.resize(1280, 750)
    window.show()
    sys.exit(app.exec_())
DATABASE.close()
